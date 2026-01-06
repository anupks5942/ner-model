import os
import re
import logging
import pdfplumber
try:
    import PyPDF2  # Fallback text extractor for PDFs pdfplumber cannot read
except ImportError:  # Keep optional to avoid hard dependency failure
    PyPDF2 = None
import spacy
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from datetime import datetime

# Initialize Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------
# 1. NLP MODEL LOADER (Prioritizes Transformer for Accuracy)
# ---------------------------------------------------------
try:
    # Try loading the Transformer model (Best Accuracy)
    nlp = spacy.load("en_core_web_trf")
    logger.info("✅ Loaded 'en_core_web_trf' (High Accuracy Model)")
except OSError:
    logger.warning("⚠️ 'en_core_web_trf' not found. Falling back to 'en_core_web_sm'.")
    logger.warning("👉 FOR MAX ACCURACY: Run 'python -m spacy download en_core_web_trf'")
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        from spacy.cli import download
        download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")

# ==========================================
# 2. ADVANCED FILE READING (The XML Hack)
# ==========================================

def read_pdf(path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.error(f"PDF read failed: {e}")

    # Fallback: pdfplumber sometimes returns empty for certain encodings; try PyPDF2 if available.
    if not text.strip() and PyPDF2 is not None:
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:  # noqa: BLE001
            logger.error(f"PyPDF2 fallback failed: {e}")

    return text.strip()

def read_docx_xml(path: str) -> str:
    """
    Parses DOCX XML directly to find text in Text Boxes, Shapes, and Tables
    that python-docx typically misses.
    """
    text_blocks = []
    try:
        with zipfile.ZipFile(path) as z:
            # List of XML files containing text (Body, Headers, Footers)
            xml_files = [f for f in z.namelist() if f.startswith("word/document") or f.startswith("word/header") or f.startswith("word/footer")]
            
            for xml_file in xml_files:
                xml_content = z.read(xml_file)
                tree = ET.fromstring(xml_content)
                
                # Iterate over all Paragraphs (<w:p>) to preserve newlines
                # We use specific namespaces usually, but simple tag matching is more robust across versions
                for node in tree.iter():
                    if node.tag.endswith('}p'): # Paragraph
                        para_text = []
                        for child in node.iter():
                            if child.tag.endswith('}t'): # Text node
                                if child.text:
                                    para_text.append(child.text)
                        if para_text:
                            text_blocks.append("".join(para_text))
                            
        return "\n".join(text_blocks)
    except Exception as e:
        logger.error(f"Deep XML Read failed: {e}")
        return ""

def read_docx_fallback(path: str) -> str:
    """Standard python-docx reader as a backup."""
    try:
        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip(): full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip(): full_text.append(para.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Standard DOCX Read failed: {e}")
        return ""

def read_any_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf": 
        return read_pdf(path)
    if ext == ".docx": 
        # Try Deep XML extraction first (for Text Boxes)
        text = read_docx_xml(path)
        # If that fails or returns empty, try standard lib
        if len(text) < 10:
            text = read_docx_fallback(path)
        return text
    
    # Text/Unknown files
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except:
        return ""

# ==========================================
# 3. EXTRACTION LOGIC
# ==========================================

def extract_email(text):
    m = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else None

def extract_phone(text):
    mobile_indicators = [
        r"(?:Mobile|Phone|Contact|Tel|Call)\s*[:\-]?\s*([+]?[0-9\s\-\(\)]{10,15})",
    ]
    for pattern in mobile_indicators:
        match = re.search(pattern, text, re.I)
        if match: return normalize_phone(match.group(1))

    text_cleaned = re.sub(r"[^\d+]", " ", text)
    candidates = re.findall(r"\+?\d{10,13}", text_cleaned)
    for c in candidates:
        return normalize_phone(c)
    return None

def normalize_phone(phone_str):
    if not phone_str: return None
    digits = re.sub(r"\D", "", phone_str)
    if len(digits) > 10: digits = digits[-10:]
    if len(digits) == 10 and digits[0] in "6789": return digits
    return None

def extract_dob(text):
    dob_indicators = [
        r"(?:DOB|Date of Birth|Birth Date|Born)[\s:\-]+(\d{1,2}[./-]\d{1,2}[./-]\d{4})",
        r"(?:DOB|Date of Birth|Birth Date|Born)[\s:\-]+(\d{1,2}\s+[A-Za-z]{3,10}[.,\s-]*\d{4})",
    ]
    for pattern in dob_indicators:
        match = re.search(pattern, text, re.I)
        if match: return match.group(1)

    general_patterns = [
        r"\b\d{1,2}[./-]\d{1,2}[./-]\d{4}\b",
        r"\b\d{1,2}\s+[A-Za-z]{3,10}[.,\s-]*\d{4}\b",
    ]
    for pattern in general_patterns:
        match = re.search(pattern, text, re.I)
        if match: return match.group(0)
    return None

def normalize_dob(dob_raw):
    if not dob_raw: return None
    dob_clean = re.sub(r"[.,]", " ", dob_raw)
    dob_clean = re.sub(r"\s+", " ", dob_clean).strip()
    formats = ["%d %B %Y", "%d %b %Y", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d.%m.%Y"]
    for fmt in formats:
        try:
            return datetime.strptime(dob_clean, fmt).date().isoformat()
        except ValueError:
            continue
    return None

def extract_gender(text):
    if re.search(r"\b(?:Male)\b", text, re.I): return "Male"
    if re.search(r"\b(?:Female)\b", text, re.I): return "Female"
    return "Not Specified"

def extract_name(text, email=None):
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    header_text = "\n".join(lines[:20])
    doc = nlp(header_text)

    # 1. EXTENDED BLOCKLIST (Contextual + Keyword)
    generic_stops = {
        # Headers
        "resume", "curriculum", "vitae", "cv", "bio", "profile", "summary",
        "education", "work", "experience", "skills", "projects", "contact",
        "mobile", "phone", "email", "address", "linkedin", "github", "link",
        "technologies", "technical", "personal", "details", "declaration",
        "objective", "career", "professional", "academic", "history",
        "reference", "references", "referee",
        
        # Job Titles
        "designer", "developer", "engineer", "manager", "consultant", "analyst",
        "head", "lead", "executive", "assistant", "associate", "intern",
        "journalist", "counselor", "plant", "graphic", "teacher", "clerk", "officer",
        
        # Organizations
        "university", "college", "school", "academy", "institute", "organization",
        "inc", "ltd", "corp", "group", "services", "solutions", "limited",
        
        # Noise
        "name", "candidate", "machine", "learning", "deep", "artificial", "intelligence", "data",
        "ca", "cpa", "bachelor", "master", "phd", "student", "graduate",
        "ma", "ba", "bs", "ms", "bsc", "msc", "bba", "mba",
        "english", "spanish", "french", "german", "language",
        
        # Relationships (Contextual Blockers)
        "mother", "father", "parent", "brother", "sister", "wife", "husband",
        "dog", "cat", "pet", "breed", "terrier"
    }

    structural_markers = {
        "street", "road", "lane", "apt", "apartment", "house", "flat", 
        "district", "state", "pincode", "zip", "nagar", "colony", "sector",
        "pradesh", "india", "complex", "building", "floor", "marg"
    }

    def is_valid_candidate(candidate_str, source_line_context=None):
        if not candidate_str: return False
        
        # Check 1: Forbidden Characters
        if re.search(r"[\d,@/()]", candidate_str): return False
            
        # Check 2: Word Count
        words = candidate_str.split()
        if len(words) < 2 or len(words) > 4: return False
            
        # Check 3: Blocklists on Candidate
        if any(w.lower() in generic_stops for w in words): return False
        if any(w.lower() in structural_markers for w in words): return False
        
        # Check 4: Context Blocklist (The Line it came from)
        if source_line_context:
            context_words = re.sub(r"[^a-zA-Z\s]", " ", source_line_context).split()
            # If the line contains "Mother" or "Reference", reject the name "Nicole"
            if any(w.lower() in generic_stops for w in context_words): return False
            if any(w.lower() in structural_markers for w in context_words): return False
            
        return True

    # STRATEGY 1: Label-Based Regex
    label_pattern = r"(?:Name|Candidate Name|Full Name)\s*[:\-]\s*([A-Za-z\s\.]+)"
    for line in lines[:20]:
        match = re.search(label_pattern, line, re.I)
        if match:
            clean_label_name = re.sub(r'\d+', '', match.group(1)).strip()
            # For label-based, skip context check since we matched the label
            if is_valid_candidate(clean_label_name):
                return clean_label_name

    candidates = []

    # STRATEGY 2: SpaCy Entity Detection
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            entity_text = ent.text.strip()
            source_line = next((line for line in lines[:20] if entity_text in line), "")
            
            # Clean full line 
            full_line_clean = re.sub(r"[^a-zA-Z\s]", " ", source_line)
            full_line_clean = re.sub(r"\s+", " ", full_line_clean).strip()
            
            # Check Full Line first, then Entity
            if is_valid_candidate(full_line_clean, source_line):
                candidates.append(full_line_clean)
            elif is_valid_candidate(re.sub(r"[^a-zA-Z\s]", "", entity_text).strip(), source_line):
                candidates.append(entity_text)

    if candidates:
        return candidates[0]

    # STRATEGY 3: Fallback (Heuristic)
    for line in lines[:10]:
        clean_text = re.sub(r"[^a-zA-Z\s]", "", line).strip()
        if is_valid_candidate(clean_text, line):
            if not re.search(r"[\d,@/]", line):
                if clean_text.istitle() or clean_text.isupper():
                    return clean_text.title()

    # STRATEGY 4: Derive from email local-part if present (e.g., john.doe@ -> John Doe)
    if email:
        local = email.split("@", 1)[0]
        local = re.sub(r"[._]+", " ", local)
        local = re.sub(r"\d+", "", local).strip()
        if local:
            candidate = " ".join(w.capitalize() for w in local.split())
            if is_valid_candidate(candidate):
                return candidate

    return None

def extract_entities(text):
    raw_dob = extract_dob(text)
    dob = normalize_dob(raw_dob) or "NA"
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "mobile": extract_phone(text),
        "dob": dob, 
        "gender": extract_gender(text)
    }

# ==========================================
# MAIN ENTRY POINT
# ==========================================
def process_resume_file(file_path: str):
    try:
        text = read_any_file(file_path)
        if not text or len(text.strip()) < 30:
            return {"error": "No readable text found"}
        return extract_entities(text)
    except Exception as e:
        logger.error(f"Processing failed: {e}")
        return {"error": str(e)}